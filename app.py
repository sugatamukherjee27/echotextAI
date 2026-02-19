import os
import re
import uuid
import logging
from datetime import datetime
from functools import wraps
from io import BytesIO

# Third-party imports
from dotenv import load_dotenv
from supabase import create_client, Client
from flask import (
    Flask, render_template, request, redirect, url_for, 
    session, flash, send_file
)
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import RequestEntityTooLarge

# AI / ML utils (Assuming these exist in your project structure)
from utils.speech_to_text import _get_whisper_model, convert_to_text
from utils.ai_summarizer import generate_output

# ---------------- CONFIGURATION ----------------
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "default-dev-key")

# File Upload Config
UPLOAD_FOLDER = os.path.join("static", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # Limit to 100MB

# AUDIO_EXTENSIONS = {"mp3", "wav", "m4a", "mp4", "avi", "mov", "mkv", "flac", "ogg"}
AUDIO_EXTENSIONS = {"mp3", "wav", "m4a", "flac", "ogg"}
VIDEO_EXTENSIONS = {"mp4", "avi", "mov", "mkv", "webm"}
IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp"}

# ---------------- SUPABASE CONNECTION ----------------
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

if not SUPABASE_URL or not SUPABASE_KEY:
    raise ValueError("SUPABASE_URL and SUPABASE_KEY must be set in the .env file")

# Initialize the client
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ---------------- UTILITIES ----------------
def allowed_file(filename, allowed_set):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_set

def get_clean_filename(output_text, extension):
    """Generates a safe filename from output text."""
    if not output_text:
        return f"document.{extension}"
    first_line = output_text.split('\n')[0].replace('*', '')
    clean_name = re.sub(r'[\\/*?:"<>|]', "", first_line).strip()
    return f"{clean_name[:50]}.{extension}"

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated_function

# ---------------- ERROR HANDLERS ----------------
@app.errorhandler(413)
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    flash("The file you uploaded is too large. Max limit is 100MB.")
    return redirect(url_for('dashboard'))

# ---------------- ROUTES ----------------

@app.route("/")
def landing():
    return render_template("land.html")

# --- AUTHENTICATION ---
@app.route("/register", methods=["GET", "POST"])
def register():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
        
    if request.method == "POST":
        username = request.form.get("username")
        first = request.form.get("firstname")
        last = request.form.get("lastname")
        email = request.form.get("email")
        password = request.form.get("password")
        date = request.form.get("date")

        if len(password) < 8:
            flash("Password must be at least 8 characters long.")
            return redirect(url_for("register"))
        
        try:
            # 1. Check if user exists
            # We select count or id where username OR email matches
            existing_user = supabase.table("users").select("id").or_(f"username.eq.{username},email.eq.{email}").execute()
            
            if existing_user.data:
                flash("Username or Email already exists")
                return redirect(url_for("register"))

            # 2. Insert new user
            user_data = {
                "username": username,
                "first_name": first,
                "last_name": last,
                "email": email,
                "password_hash": generate_password_hash(password),
                "date_of_birth": date
            }
            supabase.table("users").insert(user_data).execute()
                
            flash("Registration successful")
            return redirect(url_for("login"))
            
        except Exception as e:
            logger.error(f"Registration Error: {e}")
            flash("An error occurred during registration.")
            return redirect(url_for("register"))

    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
        
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        
        try:
            # Fetch user
            response = supabase.table("users").select("*").eq("username", username).execute()
            
            if response.data:
                user = response.data[0] # Get the first (and only) result
                
                if check_password_hash(user["password_hash"], password):
                    session.clear()
                    session.update({
                        "user_id": user["id"],
                        "username": user["username"],
                        "first_name": user["first_name"],
                        "last_name": user["last_name"],
                        "email": user["email"],
                        "date": user["date_of_birth"],
                        "profile_image": user["profile_image"]
                    })
                    return redirect(url_for("dashboard"))
            
            flash("Invalid username or password")
            
        except Exception as e:
            logger.error(f"Login Error: {e}")
            flash("Login failed due to system error.")

    return render_template("login.html")

# --- Updated Logout Route ---
@app.route("/logout")
def logout():
    session.clear()
    flash("You have been logged out.")

    response = redirect(url_for("landing"))

    # Kill browser history caching
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"

    # VERY IMPORTANT: force new navigation context
    response.headers["Clear-Site-Data"] = '"cache", "cookies", "storage"'

    return response

# --- DASHBOARD & PROCESS ---
@app.route("/dashboard")
@login_required
def dashboard():
    last_id = session.pop("last_history_id", None)
    selected = session.pop("selected", "notes") 
    result = None
    input_text = None

    if last_id:
        try:
            response = supabase.table("history").select("input_text, output_text").eq("id", last_id).execute()
            if response.data:
                record = response.data[0]
                result = record["output_text"]
                input_text = record["input_text"]
        except Exception as e:
            logger.error(f"Error fetching dashboard data: {e}")
            
    return render_template(
        "index.html",
        user_id=session["user_id"],
        result=result,
        input_text=input_text,
        selected=selected
    )

@app.route("/upload", methods=["POST"])
@login_required
def upload():  
    output_type = request.form.get("output_type", "notes")
    user_prompt = (request.form.get("user_prompt") or "").strip()
    file = request.files.get("audio_file")
    
    text = None
    input_type = None
    path = None # Track path for cleanup

    try:
        if file and file.filename:
            filename = secure_filename(file.filename)
            unique_name = f"{uuid.uuid4().hex}_{filename}"
            path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
            file.save(path) 

            try:
                text = convert_to_text(path)
            except RuntimeError as e:
                flash("AI model is starting, please retry in 10 seconds.")
                return redirect(url_for("dashboard"))
            ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""
            input_type = "video" if ext in VIDEO_EXTENSIONS else "audio"
                
        elif user_prompt:
            text = user_prompt
            input_type = "prompt"
        else:
            flash("No input provided")
            return redirect(url_for("dashboard"))
            
        if not text:
            raise ValueError("Failed to extract text from input.")

        result = generate_output(text, output_type)

        # Save to Supabase
        history_data = {
            "user_id": session["user_id"],
            "input_type": input_type,
            "input_text": text,
            "output_text": result,
            "output_type": output_type
        }
        
        response = supabase.table("history").insert(history_data).execute()
        if response.data:
            session["last_history_id"] = response.data[0]['id']
            session["selected"] = output_type
        
        return redirect(url_for("dashboard"))

    except Exception as e:
        logger.error(f"Processing error: {e}")
        flash(f"Error: {str(e)}")
        return redirect(url_for("dashboard"))
    
    finally:
        # CLEANUP: Delete the file after processing to save disk space
        if path and os.path.exists(path):
            os.remove(path)

# --- HISTORY ---
@app.route("/history")
@login_required
def history():
    query = request.args.get("q", "").strip()
    
    try:
        db_query = supabase.table("history").select("*").eq("user_id", session["user_id"]).order("created_at", desc=True)
        
        if query:
            search_pattern = f"%{query}%"
            db_query = db_query.or_(f"input_text.ilike.{search_pattern},output_text.ilike.{search_pattern}")

        response = db_query.execute()
        records = response.data

        # --- DATE FORMATTING FIX ---
        for record in records:
            raw_date = record.get("created_at")
            if raw_date:
                try:
                    # Parse the ISO string (handling potential 'Z' or offset)
                    dt_obj = datetime.fromisoformat(raw_date.replace('Z', '+00:00'))
                    # Format it to '2024-05-20' or any format you prefer
                    record["formatted_date"] = dt_obj.strftime('%Y-%m-%d')
                except Exception:
                    record["formatted_date"] = "N/A"
        # ---------------------------

    except Exception as e:
        logger.error(f"History Fetch Error: {e}")
        records = []

    return render_template("history.html", records=records, query=query)

@app.route("/history/delete/<int:history_id>", methods=["POST"])
@login_required
def delete_history_item(history_id):
    try:
        # Supabase handles permissions, but adding .eq("user_id") adds a safety layer
        supabase.table("history").delete().eq("id", history_id).eq("user_id", session["user_id"]).execute()
    except Exception as e:
        logger.error(f"Delete Error: {e}")
        flash("Could not delete item.")
    return redirect(url_for("history"))

@app.route("/history/delete_all", methods=["POST"])
@login_required
def delete_all_history():
    try:
        supabase.table("history").delete().eq("user_id", session["user_id"]).execute()
    except Exception as e:
        logger.error(f"Delete All Error: {e}")
    return redirect(url_for("history"))

# --- DOWNLOADS ---
@app.route("/history/<int:history_id>/pdf")
@login_required
def download_history_pdf(history_id):
    try:
        response = supabase.table("history").select("*").eq("id", history_id).eq("user_id", session["user_id"]).execute()
        if not response.data:
            return redirect(url_for("history"))
        record = response.data[0]
    except Exception:
        return redirect(url_for("history"))
        
    fname = get_clean_filename(record["output_text"], "pdf")
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("<b>Input:</b><br/>" + (record["input_text"] or "").replace("\n", "<br/>"), styles["Normal"]),
        Paragraph("<br/><b>Output:</b><br/>" + (record["output_text"] or "").replace("\n", "<br/>"), styles["Normal"])
    ]
    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=fname, mimetype="application/pdf")

@app.route("/history/<int:history_id>/docx")
@login_required
def download_history_docx(history_id):
    try:
        response = supabase.table("history").select("*").eq("id", history_id).eq("user_id", session["user_id"]).execute()
        if not response.data:
            return redirect(url_for("history"))
        record = response.data[0]
    except Exception:
        return redirect(url_for("history"))
        
    fname = get_clean_filename(record["output_text"], "docx")
    doc = Document()
    doc.add_heading(fname.replace(".docx", ""), level=1)
    doc.add_paragraph(record["output_text"])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=fname, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/history/<int:history_id>/txt")
@login_required
def download_history_txt(history_id):
    try:
        response = supabase.table("history").select("*").eq("id", history_id).eq("user_id", session["user_id"]).execute()
        if not response.data:
            return redirect(url_for("history"))
        record = response.data[0]
    except Exception:
        return redirect(url_for("history"))

    fname = get_clean_filename(record["output_text"], "txt")
    content = f"OUTPUT:\n{record['output_text']}"
    buffer = BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=fname, mimetype="text/plain")

# --- PROFILE SETTINGS ---
@app.route("/profile")
@login_required
def profile():
    dob = session.get("date")
    formatted_date = dob
    if dob:
        try:
            formatted_date = datetime.strptime(dob, "%Y-%m-%d").strftime("%d %B %Y")
        except ValueError:
            pass 

    return render_template(
        "profile.html",
        username=session.get("username"),
        firstname=session.get("first_name"),
        lastname=session.get("last_name"),
        email=session.get("email"),
        date=formatted_date,
        profile_image=session.get("profile_image")
    )

@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    if request.method == "POST":
        current_password = request.form.get("current_password")
        new_password = request.form.get("new_password")
        confirm_password = request.form.get("confirm_password")
        
        if len(new_password) < 8:
            flash("New password must be at least 8 characters long.")
            return redirect(url_for("change_password"))
        if new_password != confirm_password:
            flash("New passwords do not match.")
            return redirect(url_for("change_password"))
            
        try:
            response = supabase.table("users").select("password_hash").eq("id", session["user_id"]).execute()
            user = response.data[0] if response.data else None
            
            if not user or not check_password_hash(user["password_hash"], current_password):
                flash("Current password is incorrect.")
                return redirect(url_for("change_password"))
            
            supabase.table("users").update({
                "password_hash": generate_password_hash(new_password)
            }).eq("id", session["user_id"]).execute()
            
            flash("Password changed successfully.")
            return redirect(url_for("profile"))
            
        except Exception as e:
            logger.error(f"Password Change Error: {e}")
            flash("Error changing password.")

    return render_template("change_password.html")

@app.route("/upload-profile-photo", methods=["POST"])
@login_required
def upload_profile_photo():
    file = request.files.get("profile_image")
    if not file or not file.filename:
        flash("No image selected")
        return redirect(url_for("profile"))
    if not allowed_file(file.filename, IMAGE_EXTENSIONS):
        flash("Invalid image format")
        return redirect(url_for("profile"))
        
    filename = secure_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
    file.save(path)
    
    try:
        # Get old image to delete from disk
        response = supabase.table("users").select("profile_image").eq("id", session["user_id"]).execute()
        old_image = response.data[0]["profile_image"] if response.data else None
        
        # Update DB
        supabase.table("users").update({"profile_image": unique_name}).eq("id", session["user_id"]).execute()
        
        # Cleanup old file from disk
        if old_image:
            old_path = os.path.join(app.config["UPLOAD_FOLDER"], old_image)
            if os.path.exists(old_path):
                os.remove(old_path)
                
        session["profile_image"] = unique_name
        flash("Profile photo updated")
    except Exception as e:
        logger.error(f"Photo Upload Error: {e}")
        flash("Error updating profile photo")
        
    return redirect(url_for("profile"))

@app.route("/delete-profile-photo", methods=["POST"])
@login_required
def delete_profile_photo():
    try:
        response = supabase.table("users").select("profile_image").eq("id", session["user_id"]).execute()
        old_image = response.data[0]["profile_image"] if response.data else None
        
        if old_image:
            supabase.table("users").update({"profile_image": None}).eq("id", session["user_id"]).execute()
            
            path = os.path.join(app.config["UPLOAD_FOLDER"], old_image)
            if os.path.exists(path):
                os.remove(path)
                
            session["profile_image"] = None
            flash("Profile photo deleted")
    except Exception as e:
        logger.error(f"Photo Delete Error: {e}")
        
    return redirect(url_for("profile"))

@app.route("/delete-account", methods=["POST"])
@login_required
def delete_account():
    user_id = session["user_id"]
    try:
        # Get image to delete
        response = supabase.table("users").select("profile_image").eq("id", user_id).execute()
        image = response.data[0]["profile_image"] if response.data else None
        
        # Delete user
        supabase.table("users").delete().eq("id", user_id).execute()
        
        # Cleanup file
        if image:
            img_path = os.path.join(app.config["UPLOAD_FOLDER"], image)
            if os.path.exists(img_path):
                os.remove(img_path)
                
        session.clear()
        flash("Your account has been permanently deleted.")
    except Exception as e:
        logger.error(f"Account Delete Error: {e}")
        flash("Error deleting account")
        return redirect(url_for("profile"))
        
    return redirect(url_for("landing"))

# ---------------- CACHE CONTROL ----------------
@app.after_request
def add_header(response):
    """
    Stops the browser from storing sensitive dashboard snapshots in its history.
    """
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "-1"
    return response

# ---------------- ENTRY POINT ----------------
if __name__ == "__main__":
    app.run(debug=True)